import os
import pathlib
import json

import numpy as np
import pandas as pd
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import datetime
import traceback

from PyQt5.QtCore import QThread, pyqtSignal
from natsort import natsorted


class InsertTableData(QThread):  # Если требуется вставить колонтитулы
    progress = pyqtSignal(int)  # Сигнал для прогресс бара
    status = pyqtSignal(str)  # Сигнал для статус бара
    messageChanged = pyqtSignal(str, str)
    errors = pyqtSignal()

    def __init__(self, incoming_data):  # Список переданных элементов.
        QThread.__init__(self)
        self.path_file = incoming_data['path_file']
        self.start_path = incoming_data['start_path']
        self.finish_path = incoming_data['finish_path']
        self.queue = incoming_data['queue']
        self.logging = incoming_data['logging']

    def run(self):
        try:
            time_start = datetime.datetime.now()
            for_report = {'info': {'all_file': 0, 'success': 0, 'errors': 0, 'pass': 0,
                                   'time_start': str(time_start), 'work_time': ''},
                          'unloading': {'errors': False, 'text_err': []}}
            progress = 0
            self.logging.info("\n**************************************Start**************************************\n")
            self.logging.info('Начинаем заполнение документов')
            self.status.emit('Считываем значения из файла')
            self.progress.emit(progress)
            self.logging.info('Получаем список файлов в папке')
            files = [file for file in os.listdir(self.start_path) if file.endswith('.docx') and '~' not in file]
            files = natsorted(files, key=lambda y: y.rpartition(' ')[2][:-5])
            percent = 100/len(files)
            df = pd.read_csv(self.path_file, delimiter='|', encoding='ANSI', header=None, converters={0: str, 11: str})
            self.logging.info('DataFrame заполнен')
            serial_number = ''
            incoming_errors = []
            set_number = 0
            # Проверка если есть 0 в серийнике, чтобы не падала и заполняла по среднему
            err_append = False
            df[11] = [None if len(element) == 0 else element for element in df[11]]
            if df[11].isnull().any():
                self.logging.info('В серийных номерах обнаружены пропуски, заполняем по соседнему')
                for ind, (first, second) in enumerate(zip(df[0].to_numpy(), df[11].to_numpy())):
                    if ind == 0 and second is None:
                        finding_set_num = first
                        for i in range(1, 100):
                            if finding_set_num != df.iloc[i, 0]:
                                for_report['unloading']['errors'] = True
                                for_report['unloading']['text_err'].append(f'В выгрузке не заполнен серийный номер в'
                                                                           f' комплекте {first}')
                                break
                            if finding_set_num == df.iloc[i, 0] and df.iloc[i, 11]:
                                serial_number = df.iloc[i, 11]
                                break
                    if set_number != first:
                        set_number = first
                        err_append = False
                        if (not serial_number or serial_number != second) and second:
                            serial_number = second
                    if second is None:
                        if serial_number:
                            self.logging.info(f'В строке {ind + 1} обнаружен пропуск')
                            if err_append is False:
                                incoming_errors.append(f'{first} - обнаружен пропуск серийного номера,'
                                                       f' заполнен по идентичному номеру комплекта')
                                err_append = True
                            for_report['unloading']['errors'] = True
                            for_report['unloading']['text_err'].append(f'{first} - обнаружен пропуск'
                                                                       f' серийного номера,'
                                                                       f' заполнен по идентичному номеру комплекта')
                            df.loc[ind, 11] = serial_number
            # self.logging.info('Замена целочисленных серийников на текстовые с добавлением 00 в начале')
            # try:
            #     df = df.astype({11: np.int})
            #     df = df.astype({11: np.str})
            #     df[11] = ['00' + element for element in df[11]]
            # except ValueError:
            #     self.logging.info('Изменение серийников не удалось')
            #     self.status.emit(f'Ошибка преобразования серийных номеров')
            #     self.logging.info(
            #         "\n**************************************Finish**************************************\n")
            #     self.progress.emit(0)
            #     incoming_errors.append(f'Среди серийных номеров есть текстовые символы или полностью пропущенные'
            #                            f' значения в комплекте, проверьте выгрузку')
            #     self.queue.put({'errors': incoming_errors})
            #     self.errors.emit()
            #     return
            self.logging.info(f'Начинаем заполнение')
            for file in files:
                self.status.emit(f'Обработка файла {file}')
                self.logging.info(f'Обработка файла {file}')
                for_report['info']['all_file'] += 1
                for_report[file] = {'errors': False, 'text_err': [], 's/n': False, 'compound': False,
                                    'manufacturer': False, 'model': False, 'factory number': False, 'copy': False,
                                    'del': False}
                try:
                    index_serial_num = []
                    index_set_num = []
                    serial_number = False
                    set_number = file.rpartition('.')[0].rpartition(' ')[2]
                    self.logging.info(f'Первая итерация поиска номера комплекта {set_number}')
                    # Если номер комплекта слишком длинный - это серийник. Значит номера комплекта нет
                    if len(set_number) > 7:
                        serial_number = set_number.partition('.')[2] if '.' in set_number else set_number
                        self.logging.info(f'Номера комплекта нет, серийный номер - {serial_number}')
                    if serial_number:
                        set_number = False
                    else:
                        pre_serial_number = file.rpartition('.')[0].partition(' ')[2]
                        self.logging.info(f'Номер комплекта есть, ищем серийник. Попытка поиск в {pre_serial_number}')
                        if len(pre_serial_number) > 7:
                            if ' ' in pre_serial_number:
                                serial_number = pre_serial_number.partition(' ')[0]
                                serial_number = serial_number.partition('.')[2] if '.' in serial_number\
                                    else serial_number
                            else:
                                serial_number = pre_serial_number.partition('.')[2] if '.' in pre_serial_number\
                                    else pre_serial_number
                            self.logging.info(f'Серийный номер найден - {serial_number}')
                    if serial_number:
                        index_serial_num = df[df[11] == serial_number].index.tolist()
                    if set_number:
                        index_set_num = df[df[0] == set_number].index.tolist()
                    self.logging.info(f'Поиск по выгрузке индексов:\n'
                                      f'Результат по серийному номеру: {index_serial_num}\n'
                                      f'Результат по номеру комплекта: {index_set_num}')
                    # Заканчиваем с файлом если не нашли его в текстовом файле
                    if len(index_serial_num) == 0 and len(index_set_num) == 0:
                        self.logging.info(f'Для файла {file} не найдно совпадений в выгрузке')
                        for_report.pop(file)
                        progress += percent
                        for_report['info']['pass'] += 1
                        self.progress.emit(int(progress))
                        continue
                    if serial_number and set_number and (index_serial_num != index_set_num):
                        for_report['unloading']['text_error'].append(f'В выгрузке не заполнен серийный номер в'
                                                                     f' комплекте {first}')
                        self.logging.warning(f'Не совпадают индексы для серийных номеров и номера комплекта')
                        incoming_errors.append(f'Не совпадают индексы для серийных номеров и номера комплекта в'
                                               f' комплекте {set_number} и серийном номере {serial_number}')
                        for_report[file]['errors'] = True
                        for_report[file]['text_err'].append('Не совпадают индексы для серийных номеров'
                                                            ' и номера комплекта')
                        progress += percent
                        for_report['info']['errors'] += 1
                        self.progress.emit(int(progress))
                        continue
                    number_index = index_set_num if index_set_num else index_serial_num
                    self.logging.info(f'Индексы в выгрузке - {number_index}')
                    doc = docx.Document(pathlib.Path(self.start_path, file))
                    table = doc.tables[0]
                    self.logging.info(f'Открыли файл, выбрали первую таблицу')
                    try:
                        if table.cell(1, 0).text.rpartition('s/n: ')[0]:
                            table.cell(1, 0).text = table.cell(1, 0).text.rpartition('s/n: ')[0] +\
                                                    table.cell(1, 0).text.rpartition('s/n: ')[1] +\
                                                    df.iloc[number_index[0], 11]
                        else:
                            incoming_errors.append(f'{file} - в ячейке с серийным номером не '
                                                   'найден необходимый разделитель («s/n: »)')
                            for_report[file]['errors'] = True
                            for_report[file]['text_err'].append('В ячейке с серийным номером не '
                                                                'найден необходимый разделитель («s/n: »)')
                            for_report['info']['errors'] += 1
                            progress += percent
                            continue
                    except IndexError:
                        incoming_errors.append(f'{file} - в первой таблице только заголовок')
                        for_report[file]['errors'] = True
                        for_report[file]['text_err'].append(f'В первой таблице только заголовок')
                        for_report['info']['errors'] += 1
                        progress += percent
                        continue
                    table.cell(1, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table.cell(1, 0).vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in table.cell(1, 0).paragraphs[0].runs:
                        run.font.bold = True
                    for_report[file]['s/n'] = True
                    self.logging.info(f'Заполнили вторую ячейку и отформатировали её')
                    try:
                        control = {'compound': [True, 'состав'], 'manufacturer': [True, 'изготовитель'],
                                   'model': [True, 'модель'], 'factory number': [True, 'заводской номер']}
                        for index in range(len(number_index)):
                            for ind, name_dict in enumerate(control):
                                if pd.isna(df.iloc[number_index[index], ind + 3]) is False:
                                    table.cell(index + 2, ind).text = df.iloc[number_index[index], ind + 3]
                                table.cell(index + 2, ind).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                table.cell(index + 2, ind).vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                if pd.isna(df.iloc[number_index[index], ind + 3]) and control[name_dict][0]:
                                    for_report[file]['errors'] = True
                                    for_report[file]['text_err'].append(f'У элемента выгрузки отсутствует'
                                                                        f' {control[name_dict][1]}')
                                    incoming_errors.append(f'У элемента выгрузки {file} отсутствует'
                                                           f' {control[name_dict][1]}')
                                    control[name_dict][0] = False
                        for key in control:
                            for_report[file][key] = True if control[key][0] else False
                    except IndexError:
                        for_report[file]['errors'] = True
                        for_report[file]['text_err'].append('В выгрузке и таблице кол-во техники не совпадает')
                        for_report['info']['errors'] += 1
                        progress += percent
                        for key_name in ['s/n', 'compound', 'manufacturer', 'model', 'factory number']:
                            for_report[file][key_name] = False
                        continue
                    self.logging.info(f'Заполнили и отформатировали таблицу')
                    try:
                        if for_report[file]['errors'] is False:
                            doc.save(pathlib.Path(self.finish_path, file))
                            for_report[file]['copy'] = True
                            self.logging.info(f'Сохранили новый файл')
                    except PermissionError:
                        for_report[file]['errors'] = True
                        for_report[file]['text_err'].append(f'Не удалось сохранить документ {file}')
                        for_report['info']['errors'] += 1
                        progress += percent
                        for key_name in ['s/n', 'compound', 'manufacturer', 'model', 'factory number']:
                            for_report[file][key_name] = False
                        continue
                    try:
                        if for_report[file]['errors'] is False:
                            os.remove(pathlib.Path(self.start_path, file))
                            for_report[file]['del'] = True
                            self.logging.info(f'Удалили старый файл')
                    except PermissionError:
                        for_report[file]['errors'] = True
                        for_report[file]['text_err'].append(f'Не удалось удалить документ {file}')
                        for_report['info']['errors'] += 1
                        progress += percent
                except BaseException as errors:
                    incoming_errors.append(f'{file} не обработан из-за непредвиденной ошибки')
                    self.logging.error(f'Обработка файла {file} завершилась с ошибкой')
                    self.logging.error("Ошибка:\n " + str(errors) + '\n' + traceback.format_exc())
                    for_report[file]['errors'] = True
                    for_report[file]['text_err'].append('Не обработан из-за непредвиденной ошибки')
                    for_report['info']['errors'] += 1
                if for_report[file]['errors']:
                    for_report['info']['errors'] += 1
                else:
                    for_report['info']['success'] += 1
                progress += percent
                self.progress.emit(int(progress))
            if incoming_errors:
                self.logging.info("Выводим ошибки")
                self.logging.info(incoming_errors)
                self.status.emit('Готово, есть ошибки.')  # Посылаем значние если готово
                self.queue.put({'errors': incoming_errors})
                self.errors.emit()
            else:
                self.status.emit('Готово!')  # Посылаем значние если готово
            self.progress.emit(100)  # Завершаем прогресс бар
            self.logging.info("Конец программы, время работы: " + str(datetime.datetime.now() - time_start))
            for_report['info']['work_time'] = str(datetime.datetime.now() - time_start)
            self.logging.info('for_report:' + json.dumps(for_report, ensure_ascii=False))
            self.logging.info("\n**************************************Finish**************************************\n")
        except BaseException as e:  # Если ошибка
            self.status.emit('Ошибка')  # Сообщение в статус бар
            self.logging.error("Ошибка:\n " + str(e) + '\n' + traceback.format_exc())
            self.logging.info("\n**************************************Error**************************************\n")
            self.progress.emit(0)
